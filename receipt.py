from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""Test if the text is within the variable library set up to recognize commands"""


def test_text(text, var):
    for units in var:
        if text == units[0]:
            return True


"""Define what part of the library is needed and retrieve the replacement text"""


def define_text(text, var):
    for units in var:
        if text == units[0]:
            replace = units[1]
            return replace


"""Creating the receipt class """


class Receipt:

    def __init__(self, td):
        # Setting up all the settings, variable data etc. taken from ticket_data.py
        self.settings = td
        self.location_data = td.cell_data
        # Initialize the commands to set up the variables needed and create the receipt in his entirety
        self.setup_receipt()
        self.create_receipt()

    # Creating the receipt in his entirety
    def create_receipt(self):
        # Opening the template
        file = Document('Sale.docx')
        for table in file.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Retrieve the text within the cell and compare it to the library to see if it needs to be changed
                    text = cell.text
                    if test_text(text, self.location_data):
                        # Determine the text what the text needs to be replaced with
                        replacement_text = define_text(text, self.location_data)
                        # If something is empty, don't show it
                        if replacement_text == '-0' or replacement_text == 0:
                            cell.text = cell.text.replace(text, str(''))
                        else:
                            cell.text = cell.text.replace(text, str(replacement_text))
                        # Align every paragraph within the cell to its center
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        file.save('sale_test.docx')

    def setup_receipt(self):
        for items in self.location_data:
            # Set up amounts of tickets
            if items[0] == '&_ticket#_&':
                items[1] = self.settings.total_tickets
            elif items[0] == '&_base_&':
                items[1] = self.settings.base_price
            elif items[0] == '&_old_&':
                items[1] = self.settings.grandparents
            elif items[0] == '&_young_&':
                items[1] = self.settings.children
            elif items[0] == '&_babies_&':
                items[1] = self.settings.babies
            elif items[0] == '&_parking#_&':
                items[1] = self.settings.amount_parking_tickets

            # Set up prices and single ticket discounts
            elif items[0] == '&_base_&':
                items[1] = self.settings.base_price
            elif items[0] == '&_base_discount_&':
                items[1] = 0
            elif items[0] == '&_old_discount_&':
                discount = float(self.settings.base_price) - float(self.settings.price_old)
                items[1] = f'-{str(discount)}'
            elif items[0] == '&_young_discount_&':
                discount = float(self.settings.base_price) - float(self.settings.price_young)
                items[1] = f'-{str(discount)}'
            elif items[0] == '&_baby_discount_&':
                discount = float(self.settings.base_price) - float(self.settings.price_baby)
                items[1] = f'-{str(discount)}'

            # Set up the total prices and discounts based on ones purchases
            elif items[0] == '&_base_total_&':
                items[1] = float(self.settings.total_tickets) * float(self.settings.base_price)
            elif items[0] == '&_old_total_&':
                single_discount = float(self.settings.base_price) - float(self.settings.price_old)
                total_discount = single_discount * float(self.settings.grandparents)
                items[1] = f'-{str(total_discount)}'
            elif items[0] == '&_young_total_&':
                single_discount = float(self.settings.base_price) - float(self.settings.price_young)
                total_discount = single_discount * float(self.settings.children)
                items[1] = f'-{str(total_discount)}'
            elif items[0] == '&_baby_total_&':
                single_discount = float(self.settings.base_price) - float(self.settings.price_baby)
                total_discount = single_discount * float(self.settings.babies)
                items[1] = f'-{str(total_discount)}'

            # Setup up the total discounts and prices across all purchases
            elif items[0] == '&_total_discount_&':
                items[1] = f'-{str(self.settings.total_discount)}'
            elif items[0] == '&_total_&':
                total_price = float(self.settings.total_tickets) * float(self.settings.base_price)
                items[1] = total_price + float(self.settings.price_for_parking)
            elif items[0] == '&_complete_&':
                items[1] = self.settings.total_price

            # Set up parking variable, as long as its above 0
            if items[0] == '&_parking#_&':
                if self.settings.amount_parking_tickets > 0:
                    items[1] = self.settings.amount_parking_tickets
            elif items[0] == '&_parking_&':
                if self.settings.amount_parking_tickets > 0:
                    items[1] = self.settings.price_parking_ticket
            elif items[0] == '&_parking_total_&':
                if self.settings.amount_parking_tickets > 0:
                    items[1] = self.settings.price_for_parking
            elif items[0] == '&_parking_desc_&':
                if self.settings.amount_parking_tickets > 0:
                    items[1] = 'parking tickets'

            # Set up group variable, as long as group_discount = True
            if self.settings.group_discount:
                if items[0] == '&_group_&':
                    items[1] = f'-{self.settings.people_discount}'
                if items[0] == '&_group_desc_&':
                    items[1] = 'group discount'
