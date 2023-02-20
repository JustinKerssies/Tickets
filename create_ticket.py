from docx import Document
import qrcode
from datetime import date, datetime

today = date.today()
"""Proc the QR function to create a qr code and open the template and change certain paragraphs to reflect data from
    the individual tickets"""


def create_ticket(td, ticket):
    # Proc QR function
    create_qr(td, ticket)

    # Opening the template and selecting certain paragraphs
    document = Document('ticket_template.docx')
    date_para = document.paragraphs[2]
    price_para = document.paragraphs[4]
    discount_para = document.paragraphs[6]

    # Replace any string to reflect individual variables based on the ticket
    date_para.text = date_para.text.replace('&_date_&', str(today))
    price_para.text = price_para.text.replace('&_price_&', td.base_price)
    discount_para.text = discount_para.text.replace('&_discount_&', ticket.discount)

    # Save the ticket based on the ticket code
    document.save('tickets/' + str(ticket.code) + '_ticket.docx')


'''Creating a QR code based on the current time, ticket_code, age, base price for a ticket and how much age discount
    the person gets'''


def create_qr(td, ticket):
    # Set up the filename for the qr code and ask for current date and time
    filename = 'qr/' + str(ticket.code) + '.png'
    now = datetime.now()
    current_time = now.strftime("%H:%M:%S")

    # Setting up the text that the QR code needs to show and inject it into the image
    text = f'Date = {str(today)}, {current_time}\nTicket Code = {ticket.code}\nAge = {ticket.age}\n' \
           f'Base price = {td.base_price}\nAge discount = {ticket.discount}'
    img = qrcode.make(text)
    type(img)

    # Save the image based on the filename previously mentioned
    img.save(filename)
